from PyPDF2 import PdfFileMerger
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import sys
import os
import time
import comtypes.client
import datetime

# parameters ===========================================================
# where are the files: use /
filePath = 'C:/Users/Zagreb/Documents/PythonProjects/PacketMaker/'

# names of the files. Make sure to include .pdf
attachmentFileNames = [
	'DraftAndSupplementBlinded.pdf',
	'PeerPreReview1.pdf',
	'ResponseToReviewer.pdf',
	'DraftAndSupplement-Changes.pdf',
	'DraftAndSupplement-Clean.pdf'
]

authorNames = 'testAuthorNames'
paperTitle = 'testTitle'

#dates in YYYY, M, DD
receivedDates = [
	datetime.date(2019, 2, 28),
	datetime.date(2019, 3, 20),
	datetime.date(2019, 6, 13),
	datetime.date(2019, 6, 13),
	datetime.date(2019, 6, 13)
]

attachmentNames = [
	'Original blinded paper for Peer Pre-review.',
	'Peer Pre-Review',
	'Response to Reviewer',
	'Revised paper with changes noted',
	'Revised paper (clean)'
]

outputName = 'testPacket.pdf'




# end of parameters ===========================================================


fullFileNames = []
createdFileList = []
catList = []
wdFormatPDF = 17

# function definitions

# create attachment name and date
def attachmentPara(i, txt, date, doc): 
	attachmentName = 'Attachment ' + str(i) + ": "
	pTmp = doc.add_paragraph()
	pTmp.add_run(attachmentName).bold=True
	pTmp.add_run(txt)
	pTmpFormat = pTmp.paragraph_format
	pTmpFormat.space_after = Pt(0)
	
	pTmp2 = doc.add_paragraph()
	pTmp2.add_run('Date Received: ')
	pTmp2.add_run(date.strftime("%d %B, %Y"))


# function to create table of contents
def tocWriter(authors, title, dates, attachments):
	document = Document()

	document.add_heading('Alexander and Diviya Magaro Peer Pre Review Documentation Binder', level=2)
	document.add_heading('CONFIDENTIAL PEER PRE-REVIEW INFORMATION ENCLOSED', level=3)

	p0 = document.add_paragraph()

	p1 = document.add_paragraph()
	p1.add_run('Authors: ').bold = True
	p1.add_run(authors)
	p1Format = p1.paragraph_format
	p1Format.space_after = Pt(0)


	p1a = document.add_paragraph()
	p1a.add_run('Paper: ').bold = True
	p1a.add_run(title)


	for i in range(len(attachments)):
		attachmentPara(i+1, attachments[i], dates[i], document)

	document.add_page_break()

	document.save(filePath+'TableOfContents.docx')



def attachmentHeaderWriter(authors, title, attachmentDate, attachmentName, attachmentNum):
	document = Document()

	document.add_heading('ATTACHMENT ' + str(attachmentNum+1), level=1)
	document.add_heading('Alexander and Diviya Magaro Peer Pre Review Documentation Binder', level=2)

	p0 = document.add_paragraph()

	p1 = document.add_paragraph()
	p1.add_run('Authors: ').bold = True
	p1.add_run(authors)
	p1Format = p1.paragraph_format
	p1Format.space_after = Pt(0)


	p1a = document.add_paragraph()
	p1a.add_run('Paper: ').bold = True
	p1a.add_run(title)

	attachmentPara(attachmentNum+1, attachmentName, attachmentDate, document)
	tmpName = 'AttachmentHeader' + str(attachmentNum)+'.docx'

	document.save(filePath+tmpName)






def pdf_cat(fileNames, outputName):
	merger = PdfFileMerger()

	for pdf in fileNames:
		merger.append(pdf)

	merger.write(outputName)
	merger.close()

def docToPdf(inFile, outFile):
	#in_file = os.path.abspath(sys.argv[1])
	#out_file = os.path.abspath(sys.argv[2])
	
#	inFile =  inFile.replace('/','\\')

	word = comtypes.client.CreateObject('Word.Application')
	doc = word.Documents.Open(inFile)
	doc.SaveAs(outFile, FileFormat=wdFormatPDF)
	doc.Close()
	word.Quit()

if (len(attachmentFileNames) != len(receivedDates)) or (len(attachmentNames)!= len(receivedDates)) or (len(attachmentNames)!= len(attachmentFileNames)):
	print('The lists attachmentFileNames, receivedDates, and attachmentNames must be of the same length')
	time.sleep(8)
	exit()


print("Please wait...")

print("Creating pages...")
# create and add table of contents
tocWriter(authorNames, paperTitle, receivedDates, attachmentNames)
docToPdf(filePath+'TableOfContents.docx', filePath+'TableOfContents.pdf')
catList.append('TableOfContents.pdf')
createdFileList.append(filePath+'TableOfContents.docx')
createdFileList.append(filePath+'TableOfContents.pdf')

# create and add attachment headers
for i in range(len(attachmentNames)):
	fileNameRoot = 'AttachmentHeader' +str(i)
	attachmentHeaderWriter(authorNames, paperTitle, receivedDates[i], attachmentNames[i], i)
	docToPdf(filePath+fileNameRoot+'.docx', filePath+fileNameRoot+'.pdf')

	
	catList.append(fileNameRoot+'.pdf')
	catList.append(attachmentFileNames[i])

	
	createdFileList.append(fileNameRoot+'.docx')
	createdFileList.append(fileNameRoot+'.pdf')




# create paths for pdfs
for i in range(len(catList)):
	fullFileNames.append(filePath + catList[i])

print("Merging PDFs...")
# merge all pdfs
pdf_cat(fullFileNames, outputName)


# clean up
print("Cleaning up...")
for file in createdFileList:
	os.remove(file)


print("Success!")
time.sleep(8)
